from flask import Flask, render_template, request, send_file, redirect, url_for, session, flash
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
from io import BytesIO
from mysql.connector import Error
import os
import mysql.connector

# =============================================================================
# CONFIGURACIÓN DE LA APLICACIÓN
# =============================================================================

app = Flask(__name__)
app.secret_key = 'clave_secreta_segura'

# =============================================================================
# FUNCIONES DE BASE DE DATOS
# =============================================================================

def conectar():
    """Establece conexión con la base de datos MySQL"""
    return mysql.connector.connect(
        host='localhost',          
        user='root',   
        password='Universitario12#',
        database='systembd',
        port=3307 
    )

def validar_credenciales(usuario, password):
    """Valida las credenciales del usuario contra la base de datos y retorna el tipo de usuario"""
    try:
        conexion = conectar()
        cursor = conexion.cursor()
        # Seleccionar tanto la contraseña como el nombre para determinar el tipo
        sql = "SELECT Nombre, Contraseña FROM usuarios WHERE Nombre = %s"
        cursor.execute(sql, (usuario,))
        resultado = cursor.fetchone()
        cursor.close()
        conexion.close()

        if resultado:
            nombre_en_bd = resultado[0]
            contraseña_en_bd = resultado[1]
            
            # Validar contraseña
            if contraseña_en_bd == password:
                # Retornar el tipo de usuario basado en el nombre
                if nombre_en_bd == "Administrador":
                    return "administrador"
                elif nombre_en_bd == "Usuario":
                    return "usuario"
                else:
                    return "usuario"  # Por defecto, cualquier otro usuario es tipo "usuario"
            else:
                return False
        else:
            return False
    except Error as e:
        print(f"Error en la conexión: {e}")
        return False

def get_cargos():
    try:
        connection = conectar()
        print(f"Conexión establecida: {connection is not None}")
        
        if connection is None:
            print("Error: No se pudo establecer conexión con la base de datos")
            return []
        
        cursor = connection.cursor(dictionary=True)
        cursor.execute("SELECT idCargos, NombreCargo FROM cargos ORDER BY NombreCargo")
        cargos = cursor.fetchall()
        
        print(f"Se obtuvieron {len(cargos)} cargos: {cargos}")
        return cargos
        
    except Exception as e:
        print(f"Error obteniendo cargos: {e}")
        return []

# =============================================================================
# FUNCIONES AUXILIARES PARA DOCUMENTOS
# =============================================================================

def sombrear_celda(celda, color_hex="D9D9D9"):
    """Aplica color de fondo a una celda de tabla en Word"""
    tc = celda._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)
    
def aplicar_fuente_celda(celda, fuente="Calibri", tam=11):
    """Establece la fuente y el tamaño del texto en una celda"""
    for parrafo in celda.paragraphs:
        for run in parrafo.runs:
            run.font.name = fuente
            run.font.size = Pt(tam)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), fuente)

def formatear_fecha(fecha_str, formato_origen="%Y-%m-%d", formato_destino="%d-%m-%Y"):
    """Convierte una fecha del formato YYYY-MM-DD a DD-MM-YYYY"""
    try:
        return datetime.strptime(fecha_str, formato_origen).strftime(formato_destino)
    except (ValueError, TypeError):
        return fecha_str

# =============================================================================
# RUTAS DE AUTENTICACIÓN
# =============================================================================

@app.route('/login', methods=['GET', 'POST'])
def login():
    """Maneja el inicio de sesión de usuarios"""
    if request.method == 'POST':
        usuario = request.form['usuario']
        password = request.form['password']

        # Validar credenciales con la base de datos
        tipo_usuario = validar_credenciales(usuario, password)
        
        if tipo_usuario == "administrador":
            session['usuario'] = usuario
            session['tipo_usuario'] = 'administrador'
            return redirect(url_for('administradorti'))
        elif tipo_usuario == "usuario":
            session['usuario'] = usuario
            session['tipo_usuario'] = 'usuario'
            return redirect(url_for('usuarioti'))
        else:
            flash('Usuario o contraseña incorrectos', 'error')

    return render_template('login.html')

@app.route('/logout')
def logout():
    """Cierra la sesión del usuario"""
    session.clear()
    return redirect(url_for('login'))

# =============================================================================
# CONFIGURACIÓN DE DATOS Y CONSTANTES
# =============================================================================

# Campos de mantenimiento de hardware
MANTENIMIENTO_CAMPOS = [
    'danos_fisicos', 'teclado', 'raton', 'cargador', 'puertos',
    'camara', 'microfono', 'limpieza_interna', 'limpieza_externa',
    'parlantes', 'pantalla', 'bateria', 'ethernet', 'bluetooth', 'wifi'
]

# Preguntas para mantenimiento de hardware
PREGUNTAS_HARDWARE = [
    ('¿El equipo presenta daños físicos?', 'danos_fisicos'),
    ('¿El teclado funciona correctamente?', 'teclado'),
    ('¿El ratón/trackpad funciona correctamente?', 'raton'),
    ('¿El cargador se encuentra en buen estado?', 'cargador'),
    ('¿Los puertos funcionan correctamente?', 'puertos'),
    ('¿La cámara funciona correctamente?', 'camara'),
    ('¿El micrófono funciona correctamente?', 'microfono'),
    ('¿El equipo recibió limpieza interna?', 'limpieza_interna'),
    ('¿El equipo recibió limpieza externa?', 'limpieza_externa'),
    ('¿Los parlantes funcionan correctamente?', 'parlantes'),
    ('¿La pantalla funciona correctamente?', 'pantalla'),
    ('¿La batería funciona correctamente?', 'bateria'),
    ('¿La tarjeta ethernet funciona correctamente?', 'ethernet'),
    ('¿El Bluetooth funciona correctamente?', 'bluetooth'),
    ('¿La tarjeta wifi funciona correctamente?', 'wifi'),
]

# Preguntas de mantenimiento de software
PREGUNTAS_SOFTWARE = [
    ("¿Se verificaron los programas vigentes?", "programas_vigentes"),
    ("¿Se eliminaron los programas externos?", "programas_externos"),
    ("¿Se realizó la limpieza de archivos temporales?", "limpieza_temporal"),
    ("¿Se eliminaron los perfiles antiguos?", "perfiles_antiguos"),
    ("¿Se realizaron actualizaciones de Windows?", "actualizaciones_windows"),
    ("¿Se comprobó el estado del disco duro?", "estado_disco"),
    ("¿Se realizo el backup del usuario anterior?", "backup_usuario"),
    ("¿Se realizó desfragmentación del disco duro?", "desfragmentacion"),
]

# Lista de programas por área
PROGRAMAS_POR_COLUMNA = [
    ["Anydesk", "Office 365", "Cisco VPN", "PDF24", "Microsoft Teams", "Microsoft Defender", "Acrobat Reader"],
    ["Atlas", "Auditsoft", "---", "---", "---", "---", "---"],
    ["Concar", "Starsoft", "PDT", "PLAME", "---", "---", "---"],
    ["Impresoras", "Scanners", "---", "---", "---", "---", "---"],
    ["PDT", "PLAME", "PLE", "Renta Anual", "Mis declaraciones", "PDB", "---"]
]

# Texto de observaciones
TEXTO_OBSERVACIONES = (
    "Certifico que los elementos detallados en el presente documento me han sido entregados para mi cuidado "
    "y custodia con el propósito de cumplir con las tareas y asignaciones propias de mi cargo, siendo estos "
    "de mi única y exclusiva responsabilidad. Me comprometo a usar correctamente los recursos solo para los "
    "fines establecidos, y a no instalar ni permitir la instalación de software por personal ajeno al personal "
    "de TI de Forvis Mazars Perú. De igual forma me comprometo a devolver el equipo en las mismas condiciones "
    "y con los mismos accesorios que me fue entregado, cuando se me programe algún cambio de equipo o el "
    "vínculo laboral haya culminado."
)

# =============================================================================
# FUNCIONES DE PROCESAMIENTO DE DATOS
# =============================================================================

def extraer_datos_formulario(request):
    """Extrae y procesa todos los datos del formulario"""
    # Datos personales
    datos_personales = {
        'nombre': request.form['nombre'],
        'correo_usuario': request.form['correo'].strip(),
        'cargo': request.form['cargo'],
        'usuario': request.form['usuario'],
        'telefono': request.form['telefono']
    }
    datos_personales['correo'] = f"{datos_personales['correo_usuario']}@forvismazars.com"

    # Hardware
    datos_hardware = {
        'tipo': request.form['tipo'],
        'marca': request.form['marca'],
        'modelo': request.form['modelo'],
        'serial': request.form['serial'],
        'procesador': request.form['procesador'],
        'ram': request.form['ram'],
        'disco': request.form['disco'],
        'perifericos': request.form['perifericos']
    }

    # Entrega de equipo
    datos_entrega = {
        'nombre_recibe': request.form.get('nombre_recibe', ''),
        'fecha_recibe': request.form.get('fecha_recibe', ''),
        'nombre_entrega': request.form.get('nombre_entrega', ''),
        'fecha_entrega': request.form.get('fecha_entrega', '')
    }

    # Datos del equipo
    datos_equipo = {
        'fecha_compra': request.form.get('fecha_compra', ''),
        'equipo': request.form.get('equipo', ''),
        'marca_equipo': request.form.get('marca_equipo', ''),
        'hostname': request.form.get('hostname', ''),
        'modelo_equipo': request.form.get('modelo_equipo', ''),
        'detalle': request.form.get('detalle', ''),
        'serie_equipo': request.form.get('serie_equipo', ''),
        'os_equipo': request.form.get('os', ''),
        'garantia_raw': request.form.get('garantia', '')
    }

    # Historial de usuarios
    historial_usuarios = {
        'inicio': request.form.getlist('historial_inicio[]'),
        'fin': request.form.getlist('historial_fin[]'),
        'usuario': request.form.getlist('historial_usuario[]')
    }

    # Historial de eventos
    historial_eventos = {
        'fechas': request.form.getlist('evento_fecha[]'),
        'observaciones': request.form.getlist('evento_observaciones[]')
    }

    return {
        'personales': datos_personales,
        'hardware': datos_hardware,
        'entrega': datos_entrega,
        'equipo': datos_equipo,
        'historial_usuarios': historial_usuarios,
        'historial_eventos': historial_eventos
    }

def procesar_datos_mantenimiento(request):
    """Procesa los datos de mantenimiento de hardware y software"""
    mantenimiento_data = {}
    for campo in MANTENIMIENTO_CAMPOS:
        estado_str = request.form.get(f"{campo}_sn", "false")
        estado_bool = estado_str.lower() == "true"
        detalle = request.form.get(f"{campo}_detalle", "")
        
        mantenimiento_data[campo] = {
            "estado": estado_bool,
            "detalle": detalle
        }
    
    return mantenimiento_data

def formatear_fechas_datos(datos):
    """Formatea todas las fechas en los datos del formulario"""
    # Formatear fechas en datos del equipo
    datos['equipo']['fecha_compra_formateada'] = formatear_fecha(datos['equipo']['fecha_compra'])
    datos['equipo']['garantia_formateada'] = formatear_fecha(datos['equipo']['garantia_raw']) if datos['equipo']['garantia_raw'] else ""
    
    # Formatear fechas de entrega
    datos['entrega']['fecha_recibe_formateada'] = formatear_fecha(datos['entrega']['fecha_recibe'])
    datos['entrega']['fecha_entrega_formateada'] = formatear_fecha(datos['entrega']['fecha_entrega'])
    
    return datos

def generar_nombre_archivo(nombre):
    """Genera un nombre único para el archivo basado en el nombre y timestamp"""
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    return f"Acta_Entrega_{nombre.replace(' ', '_')}_{timestamp}.docx"

# =============================================================================
# RUTA PRINCIPAL DEL FORMULARIO
# =============================================================================

@app.route('/administradorti')
def administradorti():
    """Ruta para el panel de administrador"""
    # Verificar autenticación y tipo de usuario
    if 'usuario' not in session or 'tipo_usuario' not in session:
        return redirect(url_for('login'))
    
    if session['tipo_usuario'] != 'administrador':
        flash('No tienes permisos para acceder a esta página', 'error')
        return redirect(url_for('login'))
    
    return render_template('administradorti.html')

@app.route('/', methods=['GET', 'POST'])
def usuarioti():
    """Ruta principal que maneja el formulario y genera el documento"""
    # Verificar autenticación
    if 'usuario' not in session or 'tipo_usuario' not in session:
        return redirect(url_for('login'))
    
    # Verificar que sea un usuario válido (tanto administrador como usuario pueden acceder)
    if session['tipo_usuario'] not in ['administrador', 'usuario']:
        flash('No tienes permisos para acceder a esta página', 'error')
        return redirect(url_for('login'))
    
    fecha_actual = datetime.today().strftime('%Y-%m-%d')
    cargos = get_cargos()
    print(f"Cargos obtenidos: {len(cargos)}")

    if request.method == 'POST':
        # Extraer datos del formulario
        datos = extraer_datos_formulario(request)
        mantenimiento_data = procesar_datos_mantenimiento(request)
        
        # Formatear fechas
        datos = formatear_fechas_datos(datos)
        
        # Generar nombre del archivo
        nombre_archivo = generar_nombre_archivo(datos['personales']['nombre'])
        
        # Preparar datos para el documento
        datos_colaborador = {
            "Nombre:": datos['personales']['nombre'],
            "Correo:": datos['personales']['correo'],
            "Cargo:": datos['personales']['cargo'],
            "Usuario de red:": datos['personales']['usuario'],
            "Teléfono:": datos['personales']['telefono']
        }
        
        datos_equipo_doc = {
            "Fecha de compra:": datos['equipo']['fecha_compra_formateada'],
            "Equipo:": datos['equipo']['equipo'],
            "Marca:": datos['equipo']['marca_equipo'],
            "Hostname:": datos['equipo']['hostname'],
            "Modelo:": datos['equipo']['modelo_equipo'],
            "Detalle:": datos['equipo']['detalle'],
            "Serie:": datos['equipo']['serie_equipo'],
            "Sistema Operativo:": datos['equipo']['os_equipo'],
            "Garantía:": datos['equipo']['garantia_formateada'],
        }
        
        # Generar documento Word
        file_stream = generar_documento_word(
            datos_colaborador,
            datos['hardware'],
            datos_equipo_doc,
            datos['entrega'],
            datos['historial_usuarios'],
            datos['historial_eventos'],
            mantenimiento_data,
            request
        )
        
        # Enviar archivo para descarga
        return send_file(
            file_stream,
            as_attachment=True,
            download_name=nombre_archivo,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    return render_template('usuarioti.html', fecha_actual=fecha_actual, cargos=cargos)

# =============================================================================
# FUNCIONES DE GENERACIÓN DE DOCUMENTO WORD
# =============================================================================

def crear_encabezado_documento(doc):
    """Crea el encabezado del documento con el logo"""
    section = doc.sections[0]
    header = section.header
    for p in header.paragraphs:
        p.clear()
    header_paragraph = header.add_paragraph()
    header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    logo_path = "static/logo.png"
    if os.path.exists(logo_path):
        run = header_paragraph.add_run()
        run.add_picture(logo_path, width=Inches(1.2))

def crear_titulo_documento(doc):
    """Crea el título principal del documento"""
    titulo = doc.add_paragraph()
    titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = titulo.add_run("ACTA DE ENTREGA DE EQUIPOS DE CÓMPUTO")
    run.font.name = 'Calibri'
    run.bold = True
    run.font.size = Pt(14)
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
    doc.add_paragraph("")

def crear_tabla_datos_colaborador(doc, datos_colaborador):
    """Crea la tabla con los datos del colaborador"""
    tabla = doc.add_table(rows=6, cols=2)
    tabla.style = 'Table Grid'
    tabla.allow_autofit = False

    # Título de la tabla
    tabla.cell(0, 0).merge(tabla.cell(0, 1)).text = "DATOS DEL COLABORADOR"
    celda_titulo = tabla.cell(0, 0).paragraphs[0]
    celda_titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = celda_titulo.runs[0] if celda_titulo.runs else celda_titulo.add_run()
    run.bold = True
    sombrear_celda(tabla.cell(0, 0))
    aplicar_fuente_celda(tabla.cell(0, 0))

    # Llenar datos
    for i, (k, v) in enumerate(datos_colaborador.items(), start=1):
        tabla.cell(i, 0).text = k
        tabla.cell(i, 1).text = v
        tabla.cell(i, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        tabla.cell(i, 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        sombrear_celda(tabla.cell(i, 0))
        aplicar_fuente_celda(tabla.cell(i, 0))
        aplicar_fuente_celda(tabla.cell(i, 1))

    doc.add_paragraph("")

def crear_tabla_hardware(doc, datos_hardware):
    """Crea la tabla con información del hardware"""
    tabla_hw = doc.add_table(rows=5, cols=4)
    tabla_hw.style = 'Table Grid'
    tabla_hw.allow_autofit = False

    # Título
    tabla_hw.cell(0, 0).merge(tabla_hw.cell(0, 3)).text = "HARDWARE"
    celda_hw = tabla_hw.cell(0, 0).paragraphs[0]
    celda_hw.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = celda_hw.runs[0] if celda_hw.runs else celda_hw.add_run()
    run.bold = True
    sombrear_celda(tabla_hw.cell(0, 0))
    aplicar_fuente_celda(tabla_hw.cell(0, 0))

    # Primera fila de datos
    encabezados1 = ["TIPO", "MARCA", "MODELO", "SERIAL"]
    valores1 = [datos_hardware['tipo'], datos_hardware['marca'], 
                datos_hardware['modelo'], datos_hardware['serial']]
    
    for i, texto in enumerate(encabezados1):
        celda = tabla_hw.cell(1, i)
        celda.text = texto
        sombrear_celda(celda)
        aplicar_fuente_celda(celda)
        for run in celda.paragraphs[0].runs:
            run.bold = True
        tabla_hw.cell(2, i).text = valores1[i]
        aplicar_fuente_celda(tabla_hw.cell(2, i))
    
    # Segunda fila de datos
    encabezados2 = ["PROCESADOR", "MEMORIA RAM (GB)", "DISCO (GB)", "PERIFERICOS"]
    valores2 = [datos_hardware['procesador'], datos_hardware['ram'], 
                datos_hardware['disco'], datos_hardware['perifericos']]
    
    for i, texto in enumerate(encabezados2):
        celda = tabla_hw.cell(3, i)
        celda.text = texto
        sombrear_celda(celda)
        aplicar_fuente_celda(celda)
        for run in celda.paragraphs[0].runs:
            run.bold = True
        tabla_hw.cell(4, i).text = valores2[i]
        aplicar_fuente_celda(tabla_hw.cell(4, i))

    doc.add_paragraph("")

def crear_tabla_observaciones(doc):
    """Crea la tabla de observaciones"""
    tabla_obs = doc.add_table(rows=2, cols=1)
    tabla_obs.style = 'Table Grid'
    tabla_obs.allow_autofit = False

    # Título
    celda_encabezado = tabla_obs.cell(0, 0)
    celda_encabezado.text = "OBSERVACIONES"
    sombrear_celda(celda_encabezado)
    aplicar_fuente_celda(celda_encabezado)
    celda_encabezado.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    celda_encabezado.paragraphs[0].runs[0].bold = True

    # Contenido
    parrafo = tabla_obs.cell(1, 0).add_paragraph(TEXTO_OBSERVACIONES)
    parrafo.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    aplicar_fuente_celda(tabla_obs.cell(1, 0))

    doc.add_paragraph("")

def crear_tabla_entrega(doc, datos_entrega):
    """Crea la tabla de entrega de equipo"""
    tabla_firma = doc.add_table(rows=5, cols=2)
    tabla_firma.style = 'Table Grid'
    tabla_firma.allow_autofit = False

    # Título
    tabla_firma.cell(0, 0).merge(tabla_firma.cell(0, 1)).text = "ENTREGA DE EQUIPO"
    celda_titulo = tabla_firma.cell(0, 0)
    celda_titulo.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sombrear_celda(celda_titulo)
    aplicar_fuente_celda(celda_titulo)
    celda_titulo.paragraphs[0].runs[0].bold = True

    # Subtítulos
    tabla_firma.cell(1, 0).text = "RECIBE"
    tabla_firma.cell(1, 1).text = "ENTREGA"
    for col in range(2):
        celda = tabla_firma.cell(1, col)
        celda.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        sombrear_celda(celda)
        aplicar_fuente_celda(celda)
        celda.paragraphs[0].runs[0].bold = True

    # Nombres
    nombres = [datos_entrega['nombre_recibe'], datos_entrega['nombre_entrega']]
    for col, nombre_persona in enumerate(nombres):
        celda = tabla_firma.cell(2, col)
        p = celda.paragraphs[0]
        run = p.add_run(f"\nNombre: {nombre_persona}\n")
        aplicar_fuente_celda(celda)

    # Firmas
    for col in range(2):
        celda = tabla_firma.cell(3, col)
        p = celda.paragraphs[0]
        run = p.add_run("\nFirma:\n")
        aplicar_fuente_celda(celda)

    # Fechas
    fechas = [datos_entrega['fecha_recibe_formateada'], datos_entrega['fecha_entrega_formateada']]
    for col, fecha in enumerate(fechas):
        celda = tabla_firma.cell(4, col)
        p = celda.paragraphs[0]
        run = p.add_run(f"\nFecha: {fecha}\n")
        aplicar_fuente_celda(celda)

    doc.add_paragraph("")
    doc.add_paragraph("")

def crear_tabla_datos_equipo(doc, datos_equipo):
    """Crea la tabla con los datos del equipo"""
    tabla_equipo = doc.add_table(rows=10, cols=2)
    tabla_equipo.style = 'Table Grid'
    tabla_equipo.allow_autofit = False

    # Título
    tabla_equipo.cell(0, 0).merge(tabla_equipo.cell(0, 1)).text = "DATOS DEL EQUIPO"
    celda_titulo = tabla_equipo.cell(0, 0).paragraphs[0]
    celda_titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = celda_titulo.runs[0] if celda_titulo.runs else celda_titulo.add_run()
    run.bold = True
    sombrear_celda(tabla_equipo.cell(0, 0))
    aplicar_fuente_celda(tabla_equipo.cell(0, 0))

    # Datos
    for i, (k, v) in enumerate(datos_equipo.items(), start=1):
        tabla_equipo.cell(i, 0).text = k
        tabla_equipo.cell(i, 1).text = v
        sombrear_celda(tabla_equipo.cell(i, 0))
        aplicar_fuente_celda(tabla_equipo.cell(i, 0))
        aplicar_fuente_celda(tabla_equipo.cell(i, 1))

    doc.add_paragraph("")

def crear_tabla_historial_usuarios(doc, historial_usuarios):
    """Crea la tabla del historial de usuarios"""
    tabla_historial = doc.add_table(rows=1, cols=3)
    tabla_historial.style = 'Table Grid'
    tabla_historial.allow_autofit = False

    # Título
    tabla_historial.cell(0, 0).merge(tabla_historial.cell(0, 2)).text = "HISTORIAL DE USUARIOS"
    celda_hist_titulo = tabla_historial.cell(0, 0).paragraphs[0]
    celda_hist_titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = celda_hist_titulo.runs[0] if celda_hist_titulo.runs else celda_hist_titulo.add_run()
    run.bold = True
    sombrear_celda(tabla_historial.cell(0, 0))
    aplicar_fuente_celda(tabla_historial.cell(0, 0))

    # Encabezados
    encabezados_historial = ["INICIO", "FIN", "USUARIO"]
    fila_encabezado = tabla_historial.add_row().cells
    for i, texto in enumerate(encabezados_historial):
        celda = fila_encabezado[i]
        celda.text = texto
        sombrear_celda(celda)
        aplicar_fuente_celda(celda)
        for run in celda.paragraphs[0].runs:
            run.bold = True

    # Datos del historial
    filas_agregadas = 0
    for fi, ff, us in zip(historial_usuarios['inicio'], historial_usuarios['fin'], historial_usuarios['usuario']):
        if us.strip():  
            fi_formateada = formatear_fecha(fi)
            ff_formateada = "ACTUAL" if ff == "ACTUAL" else formatear_fecha(ff)

            fila = tabla_historial.add_row().cells
            fila[0].text = fi_formateada
            fila[1].text = ff_formateada
            fila[2].text = us
            for celda in fila:
                aplicar_fuente_celda(celda)
            filas_agregadas += 1

    # Completar tabla con filas vacías
    total_filas_deseadas = 8  
    filas_existentes = filas_agregadas + 2  
    filas_faltantes = total_filas_deseadas - filas_existentes
    for _ in range(filas_faltantes):
        fila = tabla_historial.add_row().cells
        for celda in fila:
            celda.text = ""
            aplicar_fuente_celda(celda)

    doc.add_paragraph("")

def crear_tabla_historial_eventos(doc, historial_eventos):
    """Crea la tabla del historial de eventos"""
    tabla_eventos = doc.add_table(rows=1, cols=2)
    tabla_eventos.style = 'Table Grid'
    tabla_eventos.allow_autofit = False

    # Título
    tabla_eventos.cell(0, 0).merge(tabla_eventos.cell(0, 1)).text = "HISTORIAL DE EVENTOS"
    celda_eventos_titulo = tabla_eventos.cell(0, 0).paragraphs[0]
    celda_eventos_titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = celda_eventos_titulo.runs[0] if celda_eventos_titulo.runs else celda_eventos_titulo.add_run()
    run.bold = True
    sombrear_celda(tabla_eventos.cell(0, 0))
    aplicar_fuente_celda(tabla_eventos.cell(0, 0))

    # Encabezados
    encabezados_eventos = ["FECHA", "OBSERVACIONES"]
    fila_encabezado = tabla_eventos.add_row().cells
    for i, texto in enumerate(encabezados_eventos):
        celda = fila_encabezado[i]
        celda.text = texto
        sombrear_celda(celda)
        aplicar_fuente_celda(celda)
        for run in celda.paragraphs[0].runs:
            run.bold = True

    # Datos de eventos
    filas_evento_agregadas = 0
    for fecha, obs in zip(historial_eventos['fechas'], historial_eventos['observaciones']):
        if obs.strip():  
            fecha_formateada = formatear_fecha(fecha)
            fila = tabla_eventos.add_row().cells
            fila[0].text = fecha_formateada
            fila[1].text = obs
            for celda in fila:
                aplicar_fuente_celda(celda)
            filas_evento_agregadas += 1

    # Completar tabla con filas vacías
    total_filas_eventos_deseadas = 8 
    filas_eventos_existentes = filas_evento_agregadas + 2 
    filas_eventos_faltantes = total_filas_eventos_deseadas - filas_eventos_existentes
    for _ in range(filas_eventos_faltantes):
        fila = tabla_eventos.add_row().cells
        for celda in fila:
            celda.text = ""
            aplicar_fuente_celda(celda)

    doc.add_paragraph("")

def crear_tabla_mantenimiento_hardware(doc, request):
    """Crea la tabla de mantenimiento de hardware"""
    tabla_hardware = doc.add_table(rows=17, cols=3)
    tabla_hardware.style = 'Table Grid'
    tabla_hardware.allow_autofit = False

    # Configurar anchos de columna
    anchos = [8, 2.25, 5]
    for col_idx, ancho in enumerate(anchos):
        for row_idx in range(len(tabla_hardware.rows)):
            tabla_hardware.cell(row_idx, col_idx).width = Cm(ancho)

    # Título
    tabla_hardware.cell(0, 0).merge(tabla_hardware.cell(0, 2)).text = "MANTENIMIENTO DE HARDWARE"
    celda_titulo = tabla_hardware.cell(0, 0).paragraphs[0]
    celda_titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = celda_titulo.runs[0] if celda_titulo.runs else celda_titulo.add_run()
    run.bold = True
    sombrear_celda(tabla_hardware.cell(0, 0))
    aplicar_fuente_celda(tabla_hardware.cell(0, 0))

    # Encabezados
    tabla_hardware.cell(1, 0).text = "ACTIVIDAD"
    tabla_hardware.cell(1, 1).text = "SI / NO"
    tabla_hardware.cell(1, 2).text = "DETALLES"

    for col in range(3):
        celda = tabla_hardware.cell(1, col)
        sombrear_celda(celda)
        aplicar_fuente_celda(celda)
        celda.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = celda.paragraphs[0].runs[0]
        run.bold = True

    # Datos
    for i, (texto_pregunta, clave) in enumerate(PREGUNTAS_HARDWARE, start=2):
        tabla_hardware.cell(i, 0).text = texto_pregunta
        sombrear_celda(tabla_hardware.cell(i, 0))
        aplicar_fuente_celda(tabla_hardware.cell(i, 0))

        valor_sn = request.form.get(f"{clave}_sn", "false")
        texto_sn = "Sí" if valor_sn.lower() == "true" else "No"
        tabla_hardware.cell(i, 1).text = texto_sn
        aplicar_fuente_celda(tabla_hardware.cell(i, 1))

        detalle = request.form.get(f"{clave}_detalle", "")
        tabla_hardware.cell(i, 2).text = detalle
        aplicar_fuente_celda(tabla_hardware.cell(i, 2))

    doc.add_paragraph("")

def crear_tabla_mantenimiento_software(doc, request):
    """Crea la tabla de mantenimiento de software"""
    tabla_software = doc.add_table(rows=10, cols=3)
    tabla_software.style = 'Table Grid'
    tabla_software.allow_autofit = False

    # Configurar anchos
    anchos = [8, 2.25, 5]
    for col_idx, ancho in enumerate(anchos):
        for row_idx in range(len(tabla_software.rows)):
            tabla_software.cell(row_idx, col_idx).width = Cm(ancho)

    # Título
    tabla_software.cell(0, 0).merge(tabla_software.cell(0, 2)).text = "MANTENIMIENTO DE SOFTWARE"
    celda_titulo = tabla_software.cell(0, 0).paragraphs[0]
    celda_titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = celda_titulo.runs[0] if celda_titulo.runs else celda_titulo.add_run()
    run.bold = True
    sombrear_celda(tabla_software.cell(0, 0))
    aplicar_fuente_celda(tabla_software.cell(0, 0))

    # Encabezados
    tabla_software.cell(1, 0).text = "ACTIVIDAD"
    tabla_software.cell(1, 1).text = "SI / NO"
    tabla_software.cell(1, 2).text = "DETALLES"

    for col in range(3):
        celda = tabla_software.cell(1, col)
        sombrear_celda(celda)
        aplicar_fuente_celda(celda)
        celda.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = celda.paragraphs[0].runs[0]
        run.bold = True

    # Datos
    for i, (pregunta, clave) in enumerate(PREGUNTAS_SOFTWARE, start=2):
        tabla_software.cell(i, 0).text = pregunta
        sombrear_celda(tabla_software.cell(i, 0))
        aplicar_fuente_celda(tabla_software.cell(i, 0))

        valor_sn = request.form.get(f"{clave}_sn", "false")
        texto_sn = "Sí" if valor_sn.lower() == "true" else "No"
        tabla_software.cell(i, 1).text = texto_sn
        aplicar_fuente_celda(tabla_software.cell(i, 1))

        detalle = request.form.get(f"{clave}_detalle", "")
        tabla_software.cell(i, 2).text = detalle
        aplicar_fuente_celda(tabla_software.cell(i, 2))

    doc.add_paragraph()

def crear_tabla_programas_area(doc):
    """Crea la tabla de programas por área"""
    # Título
    titulo = doc.add_paragraph()
    titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = titulo.add_run("LISTA DE PROGRAMAS POR ÁREA")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Calibri'
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
    
    # Tabla
    tabla_obs = doc.add_table(rows=9, cols=5)
    tabla_obs.style = 'Table Grid'
    tabla_obs.allow_autofit = False

    # Encabezado "Específico"
    celda_esp = tabla_obs.cell(0, 1).merge(tabla_obs.cell(0, 4))
    celda_esp.text = "Específico"
    parrafo = celda_esp.paragraphs[0]
    parrafo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = parrafo.runs[0] if parrafo.runs else parrafo.add_run()
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    sombrear_celda(celda_esp)
    aplicar_fuente_celda(celda_esp)

    # Encabezado "General"
    celda_general = tabla_obs.cell(1, 0).merge(tabla_obs.cell(0, 0))
    celda_general.text = "General (Todos)"
    parrafo = celda_general.paragraphs[0]
    parrafo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = parrafo.runs[0] if parrafo.runs else parrafo.add_run()
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    celda_general.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    sombrear_celda(celda_general)
    aplicar_fuente_celda(celda_general)

    # Áreas específicas
    areas = ["AUDIT", "AOS", "ADMIN", "TAX & LEGAL"]
    for col_idx, area in enumerate(areas, start=1):
        celda = tabla_obs.cell(1, col_idx)
        celda.text = area
        parrafo = celda.paragraphs[0]
        parrafo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = parrafo.runs[0] if parrafo.runs else parrafo.add_run()
        run.bold = True
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        sombrear_celda(celda)
        aplicar_fuente_celda(celda)

    # Llenar programas
    for col_idx, programas in enumerate(PROGRAMAS_POR_COLUMNA):
        for row_offset, programa in enumerate(programas, start=2):
            celda = tabla_obs.cell(row_offset, col_idx)
            celda.text = programa
            parrafo = celda.paragraphs[0]
            run = parrafo.runs[0] if parrafo.runs else parrafo.add_run()
            run.font.name = "Calibri"
            run.font.size = Pt(11)

def generar_documento_word(datos_colaborador, datos_hardware, datos_equipo, datos_entrega, 
                          historial_usuarios, historial_eventos, mantenimiento_data, request):
    """Función principal para generar el documento Word completo"""
    doc = Document()

    # Crear todas las secciones del documento
    crear_encabezado_documento(doc)
    crear_titulo_documento(doc)
    crear_tabla_datos_colaborador(doc, datos_colaborador)
    crear_tabla_hardware(doc, datos_hardware)
    crear_tabla_observaciones(doc)
    crear_tabla_entrega(doc, datos_entrega)
    crear_tabla_datos_equipo(doc, datos_equipo)
    crear_tabla_historial_usuarios(doc, historial_usuarios)
    crear_tabla_historial_eventos(doc, historial_eventos)
    crear_tabla_mantenimiento_hardware(doc, request)
    crear_tabla_mantenimiento_software(doc, request)
    crear_tabla_programas_area(doc)

    # Guardar en memoria
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    return file_stream

# =============================================================================
# PUNTO DE ENTRADA DE LA APLICACIÓN
# =============================================================================

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True)



